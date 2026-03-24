from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import docx
import io
import json
import base64
import os
import PyPDF2
import re
import google.generativeai as genai

llave_secreta = os.getenv("GEMINI_API_KEY")

if not llave_secreta:
    raise ValueError("No se encontró GEMINI_API_KEY en las variables de entorno.")

genai.configure(api_key=llave_secreta)

modelo = genai.GenerativeModel('gemini-pro')

app = FastAPI(title="Auditor Legal IA - Enterprise Final")

def leer_pdf_completo(ruta_archivo):
    texto_extraido = ""
    try:
        with open(ruta_archivo, "rb") as archivo:
            lector = PyPDF2.PdfReader(archivo)
            for pagina in lector.pages:
                texto_extraido += pagina.extract_text() + "\n"
    except Exception:
        texto_extraido = "Error al leer el documento legal."
    return texto_extraido

LEY_RGPD_COMPLETA = leer_pdf_completo("RGPD_SPAIN.pdf")
LEY_DORA_COMPLETA = leer_pdf_completo("DORA_SPAIN.pdf")

# --- MODELOS DE DATOS ---
class PeticionContrato(BaseModel):
    nombre_archivo: str
    archivo_base64: str

class PeticionPregunta(BaseModel):
    archivo_base64: str
    pregunta: str

# --- ENDPOINT 1: LA AUDITORÍA ESTRICTA ---
@app.post("/auditar-contrato")
async def auditar_contrato(peticion: PeticionContrato):
    try:
        texto_b64 = peticion.archivo_base64
        if "contentBytes" in texto_b64:
            try:
                obj = json.loads(texto_b64)
                texto_b64 = obj.get("contentBytes", texto_b64)
            except:
                pass
        
        contenido_binario = base64.b64decode(texto_b64)
        documento_io = io.BytesIO(contenido_binario)
        doc = docx.Document(documento_io)
        texto_completo = "\n".join([parrafo.text for parrafo in doc.paragraphs if parrafo.text.strip()])
        
        prompt_maestro = f"""
        Eres un auditor legal experto y MUY ESTRICTO. Evalúa el contrato contra el texto íntegro del Reglamento General de Protección de Datos (RGPD) y el Reglamento DORA.
        
        [LEY EUROPEA DE REFERENCIA RGPD]
        {LEY_RGPD_COMPLETA}
        
        [NORMATIVA BANCARIA DORA]
        {LEY_DORA_COMPLETA}
        
        [CONTRATO A ANALIZAR]
        {texto_completo}
        
        TAREAS OBLIGATORIAS:
        1. Extrae la información básica: proveedor, tipo_contrato, duracion, fecha.
        2. Evalúa EXHAUSTIVAMENTE 5 controles RGPD y 6 controles DORA.
        
        REGLAS DE FORMATO ESTRICTAS:
        - 'estado': DEBE SER EXACTAMENTE "OK", "Falta", o "Riesgo".
        - 'observacion': Máximo 12 palabras.
        - 'evidencia_textual': Inicia indicando sección entre corchetes, luego máximo 15 palabras.

        Devuelve ÚNICAMENTE un JSON con esta estructura:
        {{
          "informacion_basica": {{ "proveedor": "", "tipo_contrato": "", "duracion": "", "fecha": "" }},
          "cumplimiento_rgpd": [ {{"control": "", "estado": "", "observacion": "", "evidencia_textual": ""}} ],
          "cumplimiento_dora": [ {{"control": "", "estado": "", "observacion": "", "evidencia_textual": ""}} ]
        }}
        """

        # Auditoría sin límite estricto de palabras porque el JSON es grande
        respuesta = await modelo.generate_content_async(
            prompt_maestro,
            generation_config={"response_mime_type": "application/json"}
        )
        datos = json.loads(respuesta.text)

        lista_estados_rgpd = [item.get("estado", "OK") for item in datos.get("cumplimiento_rgpd", [])]
        lista_estados_dora = [item.get("estado", "OK") for item in datos.get("cumplimiento_dora", [])]
        todos_los_estados = lista_estados_rgpd + lista_estados_dora

        total_faltas = todos_los_estados.count("Falta")
        total_riesgos = todos_los_estados.count("Riesgo")

        if total_faltas >= 2 or total_riesgos >= 3:
            nivel = "Bajo"
            recomendacion = "Revisión legal obligatoria"
        elif total_faltas > 0 or total_riesgos > 0:
            nivel = "Medio"
            recomendacion = "Aprobable con cambios"
        else:
            nivel = "Alto"
            recomendacion = "Aprobable"

        datos["resultado_final"] = {
            "nivel_cumplimiento": nivel,
            "recomendacion": recomendacion
        }
        
        # 🔥 CORRECCIÓN 2: Limpieza de RAM
        del texto_completo
        del prompt_maestro
        
        return datos

    except Exception as e:
        # 🔥 CORRECCIÓN 3: Imprimir el error real en la pantalla negra de Render
        print(f"🔥 ERROR INTERNO EN AUDITORIA: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


# --- ENDPOINT 2: EL NUEVO CHATBOT CONVERSACIONAL (INFALIBLE) ---
@app.post("/preguntar-contrato")
async def preguntar_contrato(peticion: PeticionPregunta):
    try:
        texto_b64 = peticion.archivo_base64
        if "contentBytes" in texto_b64:
            try:
                obj = json.loads(texto_b64)
                texto_b64 = obj.get("contentBytes", texto_b64)
            except:
                pass
        
        contenido_binario = base64.b64decode(texto_b64)
        documento_io = io.BytesIO(contenido_binario)
        doc = docx.Document(documento_io)
        texto_completo = "\n".join([parrafo.text for parrafo in doc.paragraphs if parrafo.text.strip()])
        
        prompt_qa = f"""
        Eres un asistente legal experto. Responde a la pregunta basándote EXCLUSIVAMENTE en el contrato.
        
        REGLAS:
        1. NO uses información de internet.
        2. Si no está en el contrato, responde: "La información solicitada no se encuentra detallada en este contrato."
        3. Sé súper conciso. NUNCA superes las 150 palabras.
        
        [CONTRATO]
        {texto_completo}
        
        [PREGUNTA DEL USUARIO]
        {peticion.pregunta}
        """
        
        respuesta = await modelo.generate_content_async(prompt_qa)
        respuesta_cruda = respuesta.text
        
        texto_limpio = re.sub(r'[^\w\s.,;:!?()\'áéíóúÁÉÍÓÚñÑüÜ-]', '', respuesta_cruda)
        texto_limpio = texto_limpio.replace('\n', ' ').replace('\r', ' ')
        texto_limpio = " ".join(texto_limpio.split()) 
        
        datos = {"respuesta": texto_limpio}
        
        del texto_completo
        del prompt_qa
        
        return datos

    except Exception as e:
        print(f"🔥 ERROR INTERNO EN PREGUNTAS: {str(e)}")
        return {"respuesta": f"Error de procesamiento: {str(e)[:100]}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)